"""Generate Hollywood-format 5-scene and 10-scene starter scripts by genre."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

_REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = _REPO_ROOT / "docs" / "genre_starter_scripts"

LineKind = Literal[
    "title",
    "written_by",
    "writer_name",
    "draft_date",
    "page_break",
    "blank",
    "transition",
    "scene",
    "action",
    "character",
    "parenthetical",
    "dialogue",
    "end",
]


@dataclass(frozen=True)
class ScriptLine:
    """A single formatted line in a screenplay document."""

    kind: LineKind
    text: str = ""


@dataclass(frozen=True)
class Screenplay:
    """Metadata and body lines for one starter screenplay."""

    title: str
    writer_name: str
    draft_date: str
    lines: list[ScriptLine]


def _set_hollywood_page_layout(document: Document) -> None:
    """Apply standard Hollywood screenplay page margins and size."""
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)


def _apply_courier(paragraph: object, bold: bool = False) -> None:
    """Set Courier 12pt on every run in a paragraph."""
    for run in paragraph.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(12)
        run.bold = bold


def _add_screenplay_line(document: Document, line: ScriptLine) -> None:
    """Add one screenplay element with Hollywood-standard alignment."""
    if line.kind == "blank":
        document.add_paragraph()
        return

    if line.kind == "page_break":
        document.add_page_break()
        return

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    if line.kind == "title":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(line.text.upper())
        _apply_courier(paragraph)
        return

    if line.kind in {"written_by", "writer_name"}:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(line.text)
        _apply_courier(paragraph)
        return

    if line.kind == "draft_date":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.add_run(line.text)
        _apply_courier(paragraph)
        return

    if line.kind == "transition":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.add_run(line.text.upper())
        _apply_courier(paragraph)
        return

    if line.kind == "scene":
        paragraph.add_run(line.text.upper())
        _apply_courier(paragraph)
        return

    if line.kind == "action":
        paragraph.add_run(line.text)
        _apply_courier(paragraph)
        return

    if line.kind == "character":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(line.text.upper())
        _apply_courier(paragraph)
        return

    if line.kind == "parenthetical":
        paragraph.paragraph_format.left_indent = Inches(1.6)
        paragraph.paragraph_format.right_indent = Inches(2.0)
        paragraph.add_run(f"({line.text})")
        _apply_courier(paragraph)
        return

    if line.kind == "dialogue":
        paragraph.paragraph_format.left_indent = Inches(1.0)
        paragraph.paragraph_format.right_indent = Inches(1.5)
        paragraph.add_run(line.text)
        _apply_courier(paragraph)
        return

    if line.kind == "end":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(line.text.upper())
        _apply_courier(paragraph)
        return


def _screenplay_to_lines(
    title: str,
    writer_name: str,
    body: list[tuple[str, ...]],
) -> Screenplay:
    """Build a Screenplay object from title metadata and body tuples."""
    lines: list[ScriptLine] = [
        ScriptLine("title", title),
        ScriptLine("blank"),
        ScriptLine("blank"),
        ScriptLine("written_by", "Written by"),
        ScriptLine("blank"),
        ScriptLine("writer_name", writer_name),
        ScriptLine("blank"),
        ScriptLine("blank"),
        ScriptLine("blank"),
        ScriptLine("blank"),
        ScriptLine("draft_date", "June 2026"),
        ScriptLine("page_break"),
        ScriptLine("transition", "FADE IN:"),
        ScriptLine("blank"),
    ]
    for item in body:
        kind = item[0]
        text = item[1] if len(item) > 1 else ""
        lines.append(ScriptLine(kind, text))
    lines.extend(
        [
            ScriptLine("blank"),
            ScriptLine("transition", "FADE OUT."),
            ScriptLine("blank"),
            ScriptLine("end", "THE END"),
        ]
    )
    return Screenplay(
        title=title,
        writer_name=writer_name,
        draft_date="June 2026",
        lines=lines,
    )


def _sci_fi_5scene() -> Screenplay:
    """Return the sci-fi 5-scene starter screenplay."""
    body: list[tuple[str, ...]] = [
        ("scene", "INT. ORBITAL RELAY STATION - COMMAND DECK - DAY"),
        (
            "action",
            "DR. KARA VENN, 40s, monitors a wall of telemetry. A GOLD DATA CHIP "
            "rests beside her keyboard. The deep-space array pings — steady, then "
            "stutters.",
        ),
        ("character", "KARA"),
        ("dialogue", "Relay Seven, confirm packet integrity."),
        ("action", "Static washes the speakers. One line of text resolves on screen: ORIGIN UNKNOWN."),
        ("scene", "INT. LUNAR MISSION CONTROL - DAY"),
        (
            "action",
            "CAPTAIN NOLAN REED, 50s, studies a holographic star map. LT. MARTINEZ "
            "enters with a printout.",
        ),
        ("character", "MARTINEZ"),
        ("dialogue", "Venn flagged an unregistered burst from the Kuiper corridor."),
        ("character", "NOLAN"),
        ("dialogue", "Tell her to hold the raw feed. No one else sees it yet."),
        ("scene", "INT. ORBITAL RELAY STATION - SERVER BAY - NIGHT"),
        (
            "action",
            "Kara isolates the corrupted stream. She swaps the GOLD DATA CHIP into a "
            "portable reader. Decrypted fragments scroll — coordinates, a vessel ID.",
        ),
        ("character", "KARA"),
        ("dialogue", "Hull registry... that's been dark for twelve years."),
        ("scene", "EXT. LUNAR TRANSIT DOCK - NIGHT"),
        (
            "action",
            "Nolan crosses the gantry carrying a BLACK CASE. Dock workers salute as "
            "he boards a shuttle marked AURORA SIX.",
        ),
        ("character", "NOLAN"),
        ("dialogue", "Prep Aurora for an unscheduled outbound. Minimal crew."),
        ("scene", "INT. ORBITAL RELAY STATION - COMMAND DECK - DAY"),
        (
            "action",
            "Kara replays the signal. A grainy image resolves — a damaged hull, a "
            "flashing beacon. She reaches for her comm.",
        ),
        ("character", "KARA"),
        ("dialogue", "Nolan, the source is real. And it's asking for you by name."),
    ]
    return _screenplay_to_lines("SIGNAL LOSS", "Jordan Ellis", body)


def _sci_fi_10scene() -> Screenplay:
    """Return the sci-fi 10-scene starter screenplay."""
    body: list[tuple[str, ...]] = [
        ("scene", "INT. ORBITAL RELAY STATION - COMMAND DECK - DAY"),
        (
            "action",
            "KARA VENN watches the array. A GOLD DATA CHIP sits in a labeled tray: "
            "KUIPER LOG — DO NOT ERASE.",
        ),
        ("character", "KARA"),
        ("dialogue", "Relay Seven, confirm packet integrity."),
        ("action", "Static. One line resolves: ORIGIN UNKNOWN."),
        ("scene", "INT. LUNAR MISSION CONTROL - BRIEFING ROOM - DAY"),
        ("action", "NOLAN REED briefs LT. MARTINEZ and DR. SINGH."),
        ("character", "NOLAN"),
        ("dialogue", "Venn caught a ghost signal. We treat it as classified until verified."),
        ("character", "SINGH"),
        ("dialogue", "If that registry is correct, the Meridian went missing in 2014."),
        ("scene", "INT. ORBITAL RELAY STATION - SERVER BAY - NIGHT"),
        (
            "action",
            "Kara copies the burst onto the GOLD DATA CHIP. A secondary file unlocks — "
            "audio only.",
        ),
        ("character", "KARA"),
        ("dialogue", "That's not telemetry. That's a voice."),
        ("scene", "INT. CREW QUARTERS - ORBITAL STATION - NIGHT"),
        ("action", "Martinez packs a go-bag. Singh reviews medical supplies."),
        ("character", "MARTINEZ"),
        ("dialogue", "Reed wants a four-person team. We launch at 0600."),
        ("scene", "EXT. LUNAR TRANSIT DOCK - NIGHT"),
        ("action", "Nolan boards Aurora Six with a BLACK CASE. The shuttle seals."),
        ("character", "NOLAN"),
        ("dialogue", "No broadcast trail. We slip quiet."),
        ("scene", "INT. ORBITAL RELAY STATION - COMMAND DECK - DAY"),
        (
            "action",
            "Kara plays the audio — distorted, urgent: Reed, do not come alone. She "
            "freezes the waveform.",
        ),
        ("character", "KARA"),
        ("dialogue", "The voice matches Captain Hale. Meridian's last commander."),
        ("scene", "INT. AURORA SIX - COCKPIT - DAY"),
        ("action", "Nolan plots a course outward. Martinez straps in beside him."),
        ("character", "MARTINEZ"),
        ("dialogue", "Venn wants us to abort. She says the coordinates are a trap."),
        ("character", "NOLAN"),
        ("dialogue", "Then we walk in with our eyes open."),
        ("scene", "INT. ORBITAL RELAY STATION - COMM BOOTH - NIGHT"),
        ("action", "Kara opens a secure line. Singh's face appears on screen."),
        ("character", "SINGH"),
        ("dialogue", "I'm coming up to relay. If Hale is alive, he won't be human-shaped."),
        ("scene", "INT. DEEP-SPACE OBSERVATION NOOK - NIGHT"),
        (
            "action",
            "Through the viewport: a faint beacon pulses at the edge of scan range. "
            "Kara compares it to the chip data — a perfect match.",
        ),
        ("character", "KARA"),
        ("dialogue", "Nolan, the Meridian is out there. And it's still breathing."),
        ("scene", "INT. AURORA SIX - COCKPIT - NIGHT"),
        (
            "action",
            "The beacon grows on the forward display. Nolan grips the BLACK CASE, "
            "jaw set.",
        ),
        ("character", "NOLAN"),
        ("dialogue", "Aurora Six to Meridian. This is Nolan Reed. Respond."),
        ("action", "Silence — then a single word crawls across comms: FINALLY."),
    ]
    return _screenplay_to_lines("SIGNAL LOSS", "Jordan Ellis", body)


def _historical_fiction_5scene() -> Screenplay:
    """Return the historical fiction 5-scene starter screenplay."""
    body: list[tuple[str, ...]] = [
        ("scene", "INT. RESISTANCE SAFE HOUSE - LYON - NIGHT - 1943"),
        (
            "action",
            "Candlelight. MARIE DUBOIS, 30s, presses a SEALED ENVELOPE flat under "
            "a floorboard. Her brother LUC, 20s, watches the street.",
        ),
        ("character", "MARIE"),
        ("dialogue", "The courier comes at dawn. Until then, no one opens the door."),
        ("scene", "INT. GERMAN FIELD OFFICE - DAY"),
        (
            "action",
            "HAUPTMANN STRASSER reviews a sketch of Marie. A sergeant sets down a "
            "confiscated BICYCLE with a bent front wheel.",
        ),
        ("character", "STRASSER"),
        ("dialogue", "She passes messages through the market on Rue Victor Hugo."),
        ("scene", "EXT. LYON MARKET - MORNING"),
        (
            "action",
            "Marie sells herbs from a stall. ANTOINE, 50s, baker, buys thyme and "
            "slides a note into her basket.",
        ),
        ("character", "ANTOINE"),
        ("dialogue", "They searched the canal house yesterday. Be careful."),
        ("scene", "INT. RESISTANCE SAFE HOUSE - NIGHT"),
        (
            "action",
            "Luc returns mud-spattered. Marie reads Antoine's note by candle — "
            "STRASSER KNOWS THE ROUTE.",
        ),
        ("character", "LUC"),
        ("dialogue", "We move the envelope tonight. I know another path."),
        ("scene", "EXT. CANAL PATH - NIGHT"),
        (
            "action",
            "Marie and Luc pedal through fog on the repaired bicycle. Headlights "
            "appear behind them. Marie clutches the SEALED ENVELOPE inside her coat.",
        ),
        ("character", "MARIE"),
        ("dialogue", "If they stop us, you run. The names cannot die with us."),
    ]
    return _screenplay_to_lines("THE WINTER DISPATCH", "Claire Montgomery", body)


def _historical_fiction_10scene() -> Screenplay:
    """Return the historical fiction 10-scene starter screenplay."""
    body: list[tuple[str, ...]] = [
        ("scene", "INT. RESISTANCE SAFE HOUSE - LYON - NIGHT - 1943"),
        ("action", "MARIE DUBOIS hides a SEALED ENVELOPE under the floor. LUC watches the street."),
        ("character", "MARIE"),
        ("dialogue", "The courier comes at dawn. No one opens the door."),
        ("scene", "INT. GERMAN FIELD OFFICE - DAY"),
        ("action", "HAUPTMANN STRASSER studies Marie's sketch. A bent BICYCLE wheel sits on his desk."),
        ("character", "STRASSER"),
        ("dialogue", "She uses the market on Rue Victor Hugo."),
        ("scene", "EXT. LYON MARKET - MORNING"),
        ("action", "Marie sells herbs. BAKER ANTOINE passes a warning in her basket."),
        ("character", "ANTOINE"),
        ("dialogue", "They searched the canal house yesterday."),
        ("scene", "INT. SAFE HOUSE - NIGHT"),
        ("action", "Luc reads a map by candlelight. Marie loads bread and the envelope into a satchel."),
        ("character", "LUC"),
        ("dialogue", "We take the river road. Strasser won't expect it."),
        ("scene", "INT. CHURCH BASEMENT - NIGHT"),
        ("action", "FATHER REYNAUD, 60s, blesses three couriers. Marie kneels briefly."),
        ("character", "REYNAUD"),
        ("dialogue", "God keep the names you carry safer than our walls."),
        ("scene", "EXT. CANAL PATH - NIGHT"),
        ("action", "Marie and Luc ride the bicycle through fog. Distant engines grow louder."),
        ("character", "MARIE"),
        ("dialogue", "If they stop us, you run."),
        ("scene", "EXT. RIVER ROAD CHECKPOINT - NIGHT"),
        ("action", "German soldiers wave down a farm cart ahead. Luc slows."),
        ("character", "LUC"),
        ("dialogue", "We cut through the orchard. Now."),
        ("scene", "EXT. ORCHARD - NIGHT"),
        ("action", "They abandon the bicycle and run. A flashlight sweeps the trees."),
        ("character", "MARIE"),
        ("dialogue", "The envelope stays dry. That's all that matters."),
        ("scene", "INT. FARMHOUSE KITCHEN - NIGHT"),
        ("action", "MADAME CLAIRE, 70s, hides them behind a pantry wall. Boots crunch outside."),
        ("character", "CLAIRE"),
        ("dialogue", "Strasser pays for whispers. I charge nothing for silence."),
        ("scene", "EXT. FARM LANE - DAWN"),
        (
            "action",
            "Marie hands the SEALED ENVELOPE to a YOUNG PRIEST on a mule cart. The "
            "sun breaks over the hills.",
        ),
        ("character", "MARIE"),
        ("dialogue", "Tell London the list is real. Tell them we are still here."),
    ]
    return _screenplay_to_lines("THE WINTER DISPATCH", "Claire Montgomery", body)


def _comedy_5scene() -> Screenplay:
    """Return the comedy 5-scene starter screenplay."""
    body: list[tuple[str, ...]] = [
        ("scene", "INT. TUXEDO SHOP - DAY"),
        (
            "action",
            "DAVE KIM, 30s, groom-to-be, holds up two identical navy tuxes. His best "
            "man TOM and cousin PETE sit on a velvet bench.",
        ),
        ("character", "DAVE"),
        ("dialogue", "Same cut, same color. No one starts a wedding brawl over lapels."),
        ("character", "TOM"),
        ("dialogue", "Last time I wore rental shoes, I lost a toenail and a girlfriend."),
        ("scene", "INT. STRIP MALL BAR - NIGHT"),
        ("action", "The guys clink beers. A MAGNETIC GUEST BOOK sits in Dave's backpack."),
        ("character", "PETE"),
        ("dialogue", "Your fiancée said no surprises. This is a surprise with a power cord."),
        ("character", "DAVE"),
        ("dialogue", "It's a guest book that displays photos. Romantic. Technically."),
        ("scene", "INT. REHEARSAL DINNER - NIGHT"),
        (
            "action",
            "Dave's future mother-in-law DIANE inspects place cards. Dave sweats "
            "near the podium.",
        ),
        ("character", "DIANE"),
        ("dialogue", "Short speeches. No props. My daughter hates props."),
        ("character", "TOM"),
        ("parenthetical", "to Dave"),
        ("dialogue", "Hide the magnetic thing. I'm begging you."),
        ("scene", "INT. HOTEL BALLROOM - DAY"),
        (
            "action",
            "Wedding morning chaos. PETE irons a sash. The MAGNETIC GUEST BOOK "
            "accidentally clings to a metal door frame.",
        ),
        ("character", "PETE"),
        ("dialogue", "Dave, your romantic technology is stuck to the service entrance."),
        ("scene", "INT. HOTEL BALLROOM - NIGHT"),
        (
            "action",
            "Reception in full swing. Dave taps the mic. The guest book lights up — "
            "and projects every rehearsal-dinner toast, including Tom's drunk one.",
        ),
        ("character", "DAVE"),
        ("dialogue", "Before I introduce my wife, I'd like to apologize to her mother, "
            "science, and this ballroom."),
    ]
    return _screenplay_to_lines("THE GROOMSMEN'S GUIDE", "Marcus Webb", body)


def _comedy_10scene() -> Screenplay:
    """Return the comedy 10-scene starter screenplay."""
    body: list[tuple[str, ...]] = [
        ("scene", "INT. TUXEDO SHOP - DAY"),
        ("action", "DAVE KIM picks identical tuxes. TOM and PETE argue about bow ties."),
        ("character", "DAVE"),
        ("dialogue", "Same cut, same color. Peace in our time."),
        ("scene", "INT. BRIDAL SHOP LOBBY - DAY"),
        ("action", "Dave waits with a smoothie. His fiancée MAYA texts: Do not buy anything stupid."),
        ("character", "DAVE"),
        ("dialogue", "Define stupid. Technically everything is stupid if you zoom out."),
        ("scene", "INT. STRIP MALL BAR - NIGHT"),
        ("action", "The guys drink. Dave reveals the MAGNETIC GUEST BOOK."),
        ("character", "PETE"),
        ("dialogue", "Your mother-in-law said no props. This is a prop with Wi-Fi."),
        ("scene", "INT. REHEARSAL DINNER - NIGHT"),
        ("action", "DIANE gives a polished toast. Tom's knee jabs Dave under the table."),
        ("character", "DIANE"),
        ("dialogue", "We want elegance. Not gadgets. Not chaos."),
        ("scene", "INT. HOTEL SUITE - NIGHT"),
        ("action", "Dave practices his speech in the mirror. The guest book hums on the desk."),
        ("character", "TOM"),
        ("dialogue", "Unplug it. Unplug your dreams. Save the marriage."),
        ("scene", "INT. HOTEL BALLROOM - DAY"),
        ("action", "Florists everywhere. The guest book sticks to a rolling cart."),
        ("character", "PETE"),
        ("dialogue", "It's following us like a golden retriever with a screen."),
        ("scene", "INT. CHAPEL ANTEROOM - DAY"),
        ("action", "Dave adjusts his cufflinks. MAYA looks radiant, suspicious."),
        ("character", "MAYA"),
        ("dialogue", "You promised no surprises. Your face is doing surprise math."),
        ("scene", "INT. HOTEL BALLROOM - NIGHT"),
        ("action", "First dance ends. Dave approaches the mic. Tom kills the lights by mistake."),
        ("character", "TOM"),
        ("dialogue", "Wrong switch! I swear on my toenail!"),
        ("scene", "INT. HOTEL BALLROOM - NIGHT"),
        ("action", "The guest book projects photos — including Tom's rehearsal toast."),
        ("character", "DAVE"),
        ("dialogue", "Maya, I wanted tonight to be perfect. Instead, you get us. All of us."),
        ("scene", "INT. HOTEL BALLROOM - NIGHT"),
        ("action", "Diane sighs, then laughs. Maya takes the mic from Dave."),
        ("character", "MAYA"),
        ("dialogue", "He bought one stupid thing. He married the right person. I'll allow it."),
    ]
    return _screenplay_to_lines("THE GROOMSMEN'S GUIDE", "Marcus Webb", body)


def _drama_5scene() -> Screenplay:
    """Return the drama 5-scene starter screenplay."""
    body: list[tuple[str, ...]] = [
        ("scene", "INT. COUNTY FAMILY COURT - DAY"),
        (
            "action",
            "ELENA MORALES, 40s, sits beside her lawyer. Her ex-husband RICHARD "
            "fidgets with a SILVER WEDDING BAND on a chain.",
        ),
        ("character", "JUDGE"),
        ("dialogue", "Temporary custody remains with the mother pending evaluation."),
        ("scene", "INT. ELENA'S KITCHEN - NIGHT"),
        (
            "action",
            "Elena's daughter SOFIA, 12, does homework. Elena washes dishes, eyes on "
            "a stack of unpaid bills.",
        ),
        ("character", "SOFIA"),
        ("dialogue", "Dad said he'd pick me up Friday. Can he?"),
        ("character", "ELENA"),
        ("dialogue", "We'll see what the court says. Eat your rice."),
        ("scene", "EXT. COMMUNITY POOL - DAY"),
        (
            "action",
            "Richard watches Sofia swim laps from the bleachers. Coach ALMA "
            "approaches.",
        ),
        ("character", "ALMA"),
        ("dialogue", "She's fast when she's not looking over her shoulder."),
        ("character", "RICHARD"),
        ("dialogue", "Tell me what she needs. Not what I want to hear."),
        ("scene", "INT. THERAPIST OFFICE - DAY"),
        ("action", "Elena and Richard sit on opposite ends of the couch. THERAPIST GINA takes notes."),
        ("character", "GINA"),
        ("dialogue", "Sofia said she feels like luggage. What do each of you hear in that?"),
        ("character", "ELENA"),
        ("dialogue", "That we failed her in public and private."),
        ("scene", "EXT. COMMUNITY POOL - DUSK"),
        (
            "action",
            "Sofia sits between Elena and Richard on the bleachers. The SILVER BAND "
            "stays in Richard's pocket.",
        ),
        ("character", "SOFIA"),
        ("dialogue", "If I choose where to sleep, does that mean I stopped loving one of you?"),
        ("action", "Elena and Richard exchange a look — raw, ashamed, trying."),
        ("character", "ELENA"),
        ("dialogue", "No, mija. It means you're honest. We're the ones catching up."),
    ]
    return _screenplay_to_lines("THE WEIGHT OF WATER", "Ana Reyes", body)


def _drama_10scene() -> Screenplay:
    """Return the drama 10-scene starter screenplay."""
    body: list[tuple[str, ...]] = [
        ("scene", "INT. COUNTY FAMILY COURT - DAY"),
        ("action", "ELENA MORALES wins temporary custody. RICHARD holds a SILVER WEDDING BAND on a chain."),
        ("character", "JUDGE"),
        ("dialogue", "Evaluation continues. Both parents attend counseling."),
        ("scene", "INT. ELENA'S KITCHEN - NIGHT"),
        ("action", "SOFIA, 12, eats quietly. Bills pile up by the sink."),
        ("character", "SOFIA"),
        ("dialogue", "Dad said Friday. Is Friday still a maybe?"),
        ("scene", "EXT. COMMUNITY POOL - DAY"),
        ("action", "Richard watches Sofia swim. Coach ALMA joins him."),
        ("character", "ALMA"),
        ("dialogue", "She flinches when adults argue near the lane line."),
        ("scene", "INT. RICHARD'S APARTMENT - NIGHT"),
        ("action", "Richard sets up a second bedroom — posters, desk lamp, unopened boxes."),
        ("character", "RICHARD"),
        ("dialogue", "She has to feel wanted here. Not parked here."),
        ("scene", "INT. SCHOOL HALLWAY - DAY"),
        ("action", "Sofia passes her parents talking tersely by the office. She walks faster."),
        ("character", "SOFIA"),
        ("dialogue", "Stop using my school as a meeting room."),
        ("scene", "INT. THERAPIST OFFICE - DAY"),
        ("action", "GINA guides Elena and Richard through a custody schedule draft."),
        ("character", "GINA"),
        ("dialogue", "Sofia said she feels like luggage. What do you hear?"),
        ("scene", "INT. ELENA'S CAR - NIGHT"),
        ("action", "Rain on the windshield. Sofia stares at her phone."),
        ("character", "ELENA"),
        ("dialogue", "Talk to me. Not at me."),
        ("character", "SOFIA"),
        ("dialogue", "I'm tired of being the message you send each other."),
        ("scene", "EXT. COMMUNITY POOL - MORNING"),
        ("action", "Sofia trains alone. Richard arrives with coffee for Alma."),
        ("character", "RICHARD"),
        ("dialogue", "Tell me how to show up without making it about me."),
        ("scene", "INT. ELENA'S KITCHEN - NIGHT"),
        ("action", "Elena finds Sofia's journal open — a drawing of two houses connected by water."),
        ("character", "ELENA"),
        ("dialogue", "We don't fix this tonight. But we stop pretending it's fine."),
        ("scene", "EXT. COMMUNITY POOL - DUSK"),
        (
            "action",
            "Sofia between her parents on the bleachers. Richard keeps the SILVER BAND "
            "in his pocket.",
        ),
        ("character", "SOFIA"),
        ("dialogue", "If I choose where to sleep, does that mean I stopped loving one of you?"),
        ("character", "ELENA"),
        ("dialogue", "No, mija. It means you're honest. We're catching up."),
    ]
    return _screenplay_to_lines("THE WEIGHT OF WATER", "Ana Reyes", body)


def _horror_5scene() -> Screenplay:
    """Return the horror 5-scene starter screenplay."""
    body: list[tuple[str, ...]] = [
        ("scene", "INT. SUBURBAN HOUSE - BASEMENT - NIGHT"),
        (
            "action",
            "LEAH PARK, 30s, clears boxes after her father's funeral. A VHS CAMCORDER "
            "and a stack of tapes sit on a shelf.",
        ),
        ("character", "LEAH"),
        ("dialogue", "Last thing I need is home movies."),
        ("scene", "INT. BASEMENT - LATER"),
        (
            "action",
            "Curiosity wins. Leah plugs in the camcorder. A tape labeled CHRISTMAS "
            "'94 plays — family dinner, her father young, laughing.",
        ),
        ("action", "The image glitches. The same dinner repeats. Her father stares into lens."),
        ("character", "ON TAPE — FATHER"),
        ("dialogue", "Leah, don't rewind this one."),
        ("scene", "INT. BASEMENT - NIGHT"),
        ("action", "Leah checks the tape label again. CHRISTMAS '94. The timestamp reads 3:07 AM."),
        ("character", "LEAH"),
        ("dialogue", "Dad, you died on a Tuesday. Not at three in the morning."),
        ("scene", "INT. BASEMENT STAIRS - NIGHT"),
        (
            "action",
            "The basement light flickers. Footsteps creak overhead — but Leah lives "
            "alone.",
        ),
        ("character", "LEAH"),
        ("dialogue", "Hello?"),
        ("scene", "INT. BASEMENT - NIGHT"),
        (
            "action",
            "On the tiny screen, her father stands in this basement, pointing at the "
            "floor trapdoor Leah never noticed. The live trapdoor sits behind her, "
            "slightly open.",
        ),
        ("character", "ON TAPE — FATHER"),
        ("dialogue", "You were always down here. You just forgot."),
    ]
    return _screenplay_to_lines("BASEMENT TAPE", "Kevin Shaw", body)


def _horror_10scene() -> Screenplay:
    """Return the horror 10-scene starter screenplay."""
    body: list[tuple[str, ...]] = [
        ("scene", "INT. SUBURBAN HOUSE - FOYER - DAY"),
        ("action", "LEAH PARK unlocks her late father's house. Keys still on a fishhook by the door."),
        ("character", "LEAH"),
        ("dialogue", "One week to empty it. Then I never come back."),
        ("scene", "INT. BASEMENT - NIGHT"),
        ("action", "Boxes everywhere. A VHS CAMCORDER and labeled tapes on a shelf."),
        ("character", "LEAH"),
        ("dialogue", "Last thing I need is home movies."),
        ("scene", "INT. BASEMENT - LATER"),
        ("action", "She plays CHRISTMAS '94. Family dinner. Laughter. Then a glitch."),
        ("character", "ON TAPE — FATHER"),
        ("dialogue", "Leah, don't rewind this one."),
        ("scene", "INT. KITCHEN - NIGHT"),
        ("action", "Leah makes tea upstairs. The basement TV glow spills under the door."),
        ("character", "LEAH"),
        ("dialogue", "I unplugged you."),
        ("scene", "INT. BASEMENT - NIGHT"),
        ("action", "The camcorder still runs. Timestamp: 3:07 AM. Her father watches the lens."),
        ("character", "ON TAPE — FATHER"),
        ("dialogue", "You were always down here."),
        ("scene", "INT. BASEMENT STAIRS - NIGHT"),
        ("action", "Footsteps creak overhead. Leah lives alone."),
        ("character", "LEAH"),
        ("dialogue", "Hello?"),
        ("scene", "INT. BASEMENT - NIGHT"),
        ("action", "Behind Leah: a FLOOR TRAPDOOR she never noticed, slightly open."),
        ("character", "LEAH"),
        ("dialogue", "That wasn't here when I was twelve."),
        ("scene", "INT. BASEMENT - NIGHT"),
        ("action", "On tape, her father points at the same trapdoor. Live and recorded align."),
        ("character", "ON TAPE — FATHER"),
        ("dialogue", "You forgot the rest of the house."),
        ("scene", "INT. BASEMENT - NIGHT"),
        ("action", "Leah kneels by the trapdoor. Cold air exhales from below."),
        ("character", "LEAH"),
        ("dialogue", "Dad, if this is a joke, I'm selling the joke."),
        ("scene", "INT. BASEMENT - NIGHT"),
        (
            "action",
            "She lifts the door. A child's night-light glows under the floor — her "
            "childhood bedroom, preserved, waiting.",
        ),
        ("character", "LEAH"),
        ("dialogue", "I never left. Did I?"),
    ]
    return _screenplay_to_lines("BASEMENT TAPE", "Kevin Shaw", body)


def _fantasy_5scene() -> Screenplay:
    """Return the fantasy 5-scene starter screenplay."""
    body: list[tuple[str, ...]] = [
        ("scene", "EXT. MIST HARBOR - DAWN"),
        (
            "action",
            "A ferry dock fades in and out of fog. MIRA ASH, 20s, river-guide, ties "
            "a BRASS COMPASS to her belt. Old PILOT GARRON counts coins from travelers.",
        ),
        ("character", "GARRON"),
        ("dialogue", "Last crossing till the moon turns. You know the price."),
        ("character", "MIRA"),
        ("dialogue", "One memory, willingly given. Nothing stolen."),
        ("scene", "INT. FERRY - DAY"),
        (
            "action",
            "Passengers sit silent. A WOLF-HOODED STRANGER pays with a silver thimble. "
            "Mira's BRASS COMPASS spins — not north, but toward the stranger.",
        ),
        ("character", "MIRA"),
        ("dialogue", "You're not listed on the manifest."),
        ("character", "STRANGER"),
        ("dialogue", "Lists lag behind the dead."),
        ("scene", "EXT. RIVER MIDCHANNEL - DAY"),
        (
            "action",
            "The water blackens. Shapes swim beneath the hull — not fish, not boats.",
        ),
        ("character", "GARRON"),
        ("dialogue", "Hold the rail. The river tests what you refuse to release."),
        ("scene", "INT. FERRY - DAY"),
        (
            "action",
            "Mira opens her palm — a memory orb, faint image of a boy laughing. The "
            "stranger reaches. Mira pulls back.",
        ),
        ("character", "MIRA"),
        ("dialogue", "I pay my own fare. Always."),
        ("scene", "EXT. FAR SHORE - DUSK"),
        (
            "action",
            "The ferry docks at a city of lantern towers. Mira steps off, compass "
            "steady now. The stranger remains aboard, hood lowered — Mira's own face, older.",
        ),
        ("character", "OLDER MIRA"),
        ("dialogue", "Next time, don't leave the boy on the near bank."),
    ]
    return _screenplay_to_lines("THE LAST FERRY", "Lin Park", body)


def _fantasy_10scene() -> Screenplay:
    """Return the fantasy 10-scene starter screenplay."""
    body: list[tuple[str, ...]] = [
        ("scene", "EXT. MIST HARBOR - DAWN"),
        ("action", "MIRA ASH ties a BRASS COMPASS to her belt. PILOT GARRON counts fares."),
        ("character", "GARRON"),
        ("dialogue", "Last crossing till the moon turns."),
        ("scene", "EXT. NEAR SHORE VILLAGE - MORNING"),
        ("action", "Mira passes a boy, ELI, 10, selling river charms."),
        ("character", "ELI"),
        ("dialogue", "Take me across someday. I want to see the lantern city."),
        ("character", "MIRA"),
        ("dialogue", "Not till you can pay without losing what keeps you kind."),
        ("scene", "INT. FERRY - DAY"),
        ("action", "Passengers board. A WOLF-HOODED STRANGER offers a silver thimble."),
        ("character", "STRANGER"),
        ("dialogue", "Lists lag behind the dead."),
        ("scene", "EXT. RIVER MIDCHANNEL - DAY"),
        ("action", "The water blackens. Shapes swim beneath the hull."),
        ("character", "GARRON"),
        ("dialogue", "The river tests what you refuse to release."),
        ("scene", "INT. FERRY - DAY"),
        ("action", "Mira's BRASS COMPASS spins toward the stranger."),
        ("character", "MIRA"),
        ("dialogue", "You're not on the manifest."),
        ("scene", "INT. FERRY HOLD - DAY"),
        ("action", "Mira finds a crate marked ASH — inside, river charms like Eli's."),
        ("character", "MIRA"),
        ("dialogue", "Someone's been paying fares with other people's keepsakes."),
        ("scene", "EXT. RIVER MIDCHANNEL - DAY"),
        ("action", "Garron rings the bell three times. Fog thickens into walls."),
        ("character", "GARRON"),
        ("dialogue", "Choose your memory before the river chooses for you."),
        ("scene", "INT. FERRY - DAY"),
        ("action", "Mira holds a memory orb — Eli laughing. The stranger reaches."),
        ("character", "MIRA"),
        ("dialogue", "I pay my own fare. Always."),
        ("scene", "EXT. FAR SHORE - DUSK"),
        ("action", "Lantern towers rise. Passengers disembark, dazed, lighter."),
        ("character", "GARRON"),
        ("dialogue", "Some arrive. Some become the toll."),
        ("scene", "EXT. FAR SHORE DOCK - DUSK"),
        (
            "action",
            "The stranger lowers her hood — OLDER MIRA. The compass points back "
            "across the water, toward Eli's village.",
        ),
        ("character", "OLDER MIRA"),
        ("dialogue", "Next time, don't leave the boy on the near bank."),
    ]
    return _screenplay_to_lines("THE LAST FERRY", "Lin Park", body)


def _core_genre_screenplays() -> list[tuple[str, Screenplay]]:
    """Return genre slug and screenplay pairs for the six original starter genres."""
    return [
        ("scifi", _sci_fi_5scene()),
        ("scifi", _sci_fi_10scene()),
        ("historical_fiction", _historical_fiction_5scene()),
        ("historical_fiction", _historical_fiction_10scene()),
        ("comedy", _comedy_5scene()),
        ("comedy", _comedy_10scene()),
        ("drama", _drama_5scene()),
        ("drama", _drama_10scene()),
        ("horror", _horror_5scene()),
        ("horror", _horror_10scene()),
        ("fantasy", _fantasy_5scene()),
        ("fantasy", _fantasy_10scene()),
    ]


def _all_screenplays() -> list[tuple[str, Screenplay]]:
    """Return genre slug and screenplay pairs for all twenty starter genres."""
    from additional_genre_screenplays import additional_genre_screenplays

    return _core_genre_screenplays() + additional_genre_screenplays()


def _scene_count(screenplay: Screenplay) -> int:
    """Count INT./EXT. scene headings in a screenplay."""
    return sum(1 for line in screenplay.lines if line.kind == "scene")


_TITLE_PAGE_KINDS = frozenset(
    {"title", "written_by", "writer_name", "draft_date", "page_break"}
)


def build_screenplay_fountain(screenplay: Screenplay, output_path: Path) -> Path:
    """Write one screenplay to standard Fountain format.

    Args:
        screenplay: Screenplay content and metadata.
        output_path: Destination ``.fountain`` path.

    Returns:
        Resolved path to the written file.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    header = (
        f"Title: {screenplay.title}\n"
        f"Author: {screenplay.writer_name}\n"
        f"Draft date: {screenplay.draft_date}\n\n"
    )
    body_parts: list[str] = []
    started = False
    for line in screenplay.lines:
        if line.kind in _TITLE_PAGE_KINDS:
            continue
        if not started:
            if line.kind == "transition" and line.text.upper().startswith("FADE IN"):
                started = True
            else:
                continue
        if line.kind == "blank":
            body_parts.append("")
        elif line.kind == "transition":
            body_parts.append(line.text.upper())
        elif line.kind == "scene":
            body_parts.extend(["", line.text.upper(), ""])
        elif line.kind == "action":
            body_parts.append(line.text)
        elif line.kind == "character":
            body_parts.append(line.text.upper())
        elif line.kind == "parenthetical":
            body_parts.append(f"({line.text})")
        elif line.kind == "dialogue":
            body_parts.append(line.text)
        elif line.kind == "end":
            body_parts.extend(["", line.text.upper()])
    body_text = "\n".join(body_parts).strip() + "\n"
    output_path.write_text(header + body_text, encoding="utf-8")
    return output_path.resolve()


def build_screenplay_docx(screenplay: Screenplay, output_path: Path) -> Path:
    """Write one screenplay to a Hollywood-formatted Word document.

    Args:
        screenplay: Screenplay content and metadata.
        output_path: Destination ``.docx`` path.

    Returns:
        Resolved path to the written document.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _set_hollywood_page_layout(document)
    for line in screenplay.lines:
        _add_screenplay_line(document, line)
    document.save(str(output_path))
    return output_path.resolve()


def build_all_genre_scripts(output_dir: Path = OUTPUT_DIR) -> list[Path]:
    """Generate every genre starter script as Word and Fountain files.

    Writes twenty 5-scene and twenty 10-scene screenplays across distinct genres.

    Args:
        output_dir: Directory where screenplay files are written.

    Returns:
        List of paths to written documents.
    """
    written: list[Path] = []
    for genre_slug, screenplay in _all_screenplays():
        scene_count = _scene_count(screenplay)
        base_name = f"{genre_slug}_starter_{scene_count}scene"
        written.append(build_screenplay_docx(screenplay, output_dir / f"{base_name}.docx"))
        written.append(build_screenplay_fountain(screenplay, output_dir / f"{base_name}.fountain"))
    return written


def main() -> None:
    """CLI entry point."""
    paths = build_all_genre_scripts()
    docx_count = sum(1 for path in paths if path.suffix == ".docx")
    fountain_count = sum(1 for path in paths if path.suffix == ".fountain")
    print(f"Wrote {docx_count} Word and {fountain_count} Fountain scripts to {OUTPUT_DIR}:")
    for path in sorted(paths, key=lambda item: item.name):
        print(f"  - {path.name}")


if __name__ == "__main__":
    main()
