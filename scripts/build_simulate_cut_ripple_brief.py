"""Build a short plain-English comparison brief for cut, edit, and orphans."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

_REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = (
    _REPO_ROOT
    / "docs"
    / "CUT_EDIT_ORPHAN_PLAIN_ENGLISH_TABLES.docx"
)


def _set_document_styles(document: Document) -> None:
    """Apply base font settings to the document's Normal style."""
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _add_title(document: Document, text: str) -> None:
    """Add a centered document title."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x11, 0x32, 0x4D)


def _add_subtitle(document: Document, text: str) -> None:
    """Add a centered subtitle under the title."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x55, 0x66)


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    """Add a section heading."""
    document.add_heading(text, level=level)


def _add_para(document: Document, text: str) -> None:
    """Add a normal paragraph."""
    document.add_paragraph(text)


def _add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Insert a simple table with a header row."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_data in rows:
        row = table.add_row().cells
        for index, value in enumerate(row_data):
            row[index].text = value
    document.add_paragraph()


def build_document() -> Document:
    """Assemble the short cut / edit / orphan comparison brief."""
    document = Document()
    _set_document_styles(document)

    _add_title(document, "Simulate Cut, Simulate Edit, and Orphan Scenes")
    _add_subtitle(document, "Short plain-English reference — what each one checks")

    # ------------------------------------------------------------------
    _add_heading(document, "1. Side-by-side overview", level=1)
    _add_table(
        document,
        ["Feature", "Question it answers", "Does it change the script?"],
        [
            [
                "Simulate cut",
                "If I remove this whole scene, which later scenes lose a setup?",
                "No — preview only",
            ],
            [
                "Simulate edit",
                "If I rewrite this scene, what story links change?",
                "No — preview only",
            ],
            [
                "Orphan scenes",
                "Which scenes sit loosely and barely connect to the rest of the story?",
                "No — report only",
            ],
        ],
    )

    # ------------------------------------------------------------------
    _add_heading(document, "2. Simulate cut — what is monitored", level=1)
    _add_para(
        document,
        "You pick one scene to remove. ScriptLens checks whether later scenes "
        "still have what they need.",
    )
    _add_table(
        document,
        ["Element watched", "What it means", "Used in cut?"],
        [
            ["People", "Same character appears again later", "Yes — strong signal"],
            [
                "Things / props",
                "Same object is introduced then reused later",
                "Yes — plant and payoff",
            ],
            [
                "Places",
                "Story returns to the same location",
                "Yes — weaker signal",
            ],
            [
                "Story facts",
                "Clear established facts (injury, ownership, date, relationship)",
                "Yes",
            ],
            [
                "Backward dialogue",
                "Lines like “after what you did” or “since that night”",
                "Yes",
            ],
            [
                "Theme / emotion / pacing",
                "Whether the scene feels important artistically",
                "No",
            ],
        ],
    )
    _add_table(
        document,
        ["What happens", "Result shown to writer"],
        [
            [
                "Later scene uniquely depends on the cut scene",
                "Warned as impacted",
            ],
            [
                "Later scene still gets the same setup from an earlier scene",
                "Not warned (treated as safe)",
            ],
            [
                "No later scene depends on it",
                "Safe to cut",
            ],
        ],
    )
    _add_para(
        document,
        "Tiny example: Scene 1 plants a revolver, Scene 2 carries it, "
        "Scene 3 fires it. Cutting Scene 1 warns Scene 2 and 3. Cutting "
        "Scene 2 may be called safe if Scene 1 already introduced the gun.",
    )

    # ------------------------------------------------------------------
    _add_heading(document, "3. Simulate edit — what is monitored", level=1)
    _add_para(
        document,
        "You rewrite one scene (for example, remove the revolver). "
        "ScriptLens re-reads the script with that change and compares "
        "before vs after.",
    )
    _add_table(
        document,
        ["Element watched", "What it means", "Used in edit?"],
        [
            ["People", "Characters added, removed, or renamed in the rewrite", "Yes"],
            ["Things / props", "Objects added or removed in the rewrite", "Yes"],
            ["Places", "Location changes in the rewrite", "Yes"],
            ["Story facts", "Facts that appear or disappear after the rewrite", "Yes"],
            [
                "Story links",
                "Connections from this scene to later scenes that appear or vanish",
                "Yes — main focus",
            ],
            [
                "Orphan count",
                "Whether the rewrite creates or clears loose scenes",
                "Yes",
            ],
            [
                "Artistic quality of the rewrite",
                "Whether the new dialogue is “better”",
                "No",
            ],
        ],
    )
    _add_table(
        document,
        ["What the writer is shown", "Meaning"],
        [
            ["Links removed", "Setups that used to reach later scenes are gone"],
            ["Links added", "New setups now reach later scenes"],
            ["Links changed", "An existing link is stronger/weaker or differently explained"],
            ["Later scenes at risk", "Scenes that lose a setup from the edited scene"],
            ["Orphan change", "Loose-scene count went up, down, or stayed the same"],
        ],
    )
    _add_para(
        document,
        "Tiny example: rewrite Scene 1 so the table is empty (no revolver). "
        "ScriptLens reports the link to later gun scenes is gone, and those "
        "later scenes are at risk.",
    )

    # ------------------------------------------------------------------
    _add_heading(document, "4. Orphan scenes — what makes a scene orphan", level=1)
    _add_para(
        document,
        "An orphan is a scene that does not meaningfully join the main story "
        "thread. It feels cuttable or under-tied.",
    )
    _add_table(
        document,
        ["What is considered", "How it helps decide"],
        [
            [
                "Shared people",
                "Do the same characters appear in other scenes?",
            ],
            [
                "Shared places",
                "Does the scene share or continue a location with nearby scenes?",
            ],
            [
                "Shared things",
                "Do props or wardrobe from this scene show up elsewhere?",
            ],
            [
                "Similar meaning / topic",
                "Does the scene feel thematically tied to other scenes?",
            ],
            [
                "Opening scenes",
                "Very early scenes can be treated as prologue, not orphans",
            ],
            [
                "Montage / special blocks",
                "Grouped shot sequences are usually not flagged as orphans",
            ],
            [
                "Forward ties only",
                "If a scene has no earlier tie but clearly feeds later scenes, "
                "it is not treated as a hard orphan",
            ],
        ],
    )
    _add_table(
        document,
        ["Orphan type", "Simple meaning"],
        [
            [
                "Hard orphan",
                "Almost no useful tie into or out of the main story",
            ],
            [
                "Small loose chain",
                "A tiny group of scenes linked to each other, but not to the main story",
            ],
            [
                "Not an orphan",
                "Tied to the main story by people, places, things, or meaning — "
                "or protected as prologue/montage",
            ],
        ],
    )
    _add_para(
        document,
        "Tiny example: a motorcycle idles in an alley, and no later scene "
        "mentions that bike or continues that beat → likely orphan. "
        "A father backstory scene that still clearly feeds later June scenes "
        "→ usually not a hard orphan.",
    )

    # ------------------------------------------------------------------
    _add_heading(document, "5. Quick comparison card", level=1)
    _add_table(
        document,
        ["", "Simulate cut", "Simulate edit", "Orphan check"],
        [
            [
                "Trigger",
                "Remove one whole scene",
                "Rewrite one scene’s text",
                "Scan the whole script",
            ],
            [
                "Main focus",
                "Later scenes that lose a setup",
                "Links gained/lost after the rewrite",
                "Scenes that barely connect",
            ],
            [
                "People / things / places",
                "Yes",
                "Yes",
                "Yes",
            ],
            [
                "Story facts",
                "Yes",
                "Yes",
                "Less central — more about overall ties",
            ],
            [
                "Meaning / topic similarity",
                "No",
                "No",
                "Yes",
            ],
            [
                "Asks “is it artistically worth keeping?”",
                "No",
                "No",
                "No",
            ],
        ],
    )

    _add_para(
        document,
        "Bottom line: Cut asks what breaks if a scene disappears. "
        "Edit asks what changes if a scene is rewritten. "
        "Orphan asks which scenes are loosely attached to the story.",
    )

    return document


def main() -> None:
    """Build and save the Word brief."""
    document = build_document()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
