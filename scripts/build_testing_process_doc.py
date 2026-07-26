"""Build a Word document describing ScriptLens conventional and unconventional testing."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

_REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = _REPO_ROOT / "docs" / "SCRIPTLENS_TESTING_PROCESS.docx"


def _set_styles(document: Document) -> None:
    """Apply base font settings to the Normal style."""
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _title(document: Document, text: str) -> None:
    """Add a centered document title."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x11, 0x32, 0x4D)


def _subtitle(document: Document, text: str) -> None:
    """Add a centered italic subtitle."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def _heading(document: Document, text: str, level: int = 1) -> None:
    """Add a section heading at the given level."""
    document.add_heading(text, level=level)


def _para(document: Document, text: str) -> None:
    """Add a body paragraph."""
    document.add_paragraph(text)


def _bullet(document: Document, text: str) -> None:
    """Add a first-level bullet."""
    document.add_paragraph(text, style="List Bullet")


def _bullet2(document: Document, text: str) -> None:
    """Add a second-level bullet."""
    document.add_paragraph(text, style="List Bullet 2")


def _number(document: Document, text: str) -> None:
    """Add a numbered list item."""
    document.add_paragraph(text, style="List Number")


def _table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Add a bordered table with a bold header row."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cells[index].text = value
    document.add_paragraph()


def _build_overview(document: Document) -> None:
    """Add the title, context, and scope sections."""
    _title(document, "ScriptLens — Testing Process")
    _subtitle(
        document,
        "Conventional and unconventional testing methods for the v3 structure engine",
    )

    _heading(document, "1. Context and Scope")
    _para(
        document,
        "Plot-contradiction detection is out of the v3 product scope. Testing "
        "therefore focuses on the structure engine: scene parsing, the scene "
        "dependency graph, orphan-scene detection, simulate cut / edit, scene "
        "function impact (setup and payoff), and the draft workflow.",
    )
    _para(
        document,
        "Standard software tests check syntax, HTTP status codes, and database "
        "integrity. ScriptLens operates on story logic: it translates fuzzy "
        "natural language (Fountain / PDF text) into a deterministic directed "
        "acyclic graph. Conventional tests alone miss semantic drift, structural "
        "over-sensitivity, and false-positive cascades, so the process below "
        "pairs a conventional backbone with targeted unconventional methods.",
    )
    _para(document, "In scope for testing:")
    _bullet(document, "Fountain / PDF ingest and scene parsing")
    _bullet(document, "Scene dependency graph (continuity and causal edges)")
    _bullet(document, "Orphan-scene detection (hard orphans and loose chains)")
    _bullet(document, "Simulate cut (delete impact) and simulate edit (edge delta)")
    _bullet(document, "Scene function impact (plant / payoff / setup roles)")
    _bullet(document, "Draft workflow (delete, apply edit, undo, export) and the API")


def _build_conventional(document: Document) -> None:
    """Add the conventional testing methods section."""
    _heading(document, "2. Conventional Testing (the backbone)")
    _para(
        document,
        "These provide deterministic, repeatable coverage and form the CI gate.",
    )

    _heading(document, "2.1 Golden-file structure corpus", level=2)
    _para(
        document,
        "Hand-authored micro-scripts (5–10 scenes), each paired with a YAML "
        "ground-truth file that asserts the expected structural outputs.",
    )
    _bullet(document, "Expected orphan scenes (with hard vs. loose-chain type)")
    _bullet(document, "Expected dependency edges (or edge-count bounds)")
    _bullet(document, "Expected simulate-cut impacted scenes and risk tier")
    _bullet(document, "Expected simulate-edit edge delta (added / removed / changed)")
    _bullet(document, "Expected scene-function roles (plant, payoff, setup)")
    _para(
        document,
        "Scored with precision / recall per capability, mirroring the existing "
        "baseline scorer. This is the direct successor to the retired "
        "planted-contradiction corpus.",
    )

    _heading(document, "2.2 Wire up existing but unused ground truth", level=2)
    _para(
        document,
        "The demo ground truth already defines expected_orphans and "
        "expected_simulate_edit, but only simulate-delete is evaluated today. "
        "Wiring the remaining sections is free coverage.",
    )

    _heading(document, "2.3 Unit and regression tests", level=2)
    _bullet(document, "Per-detector unit tests for parser, graph, and orphan logic")
    _bullet(document, "Golden-output regression tests on stable fixtures")
    _bullet(document, "Deterministic seeds so semantic embeddings stay reproducible")

    _heading(document, "2.4 False-positive (clean) corpus", level=2)
    _para(
        document,
        "Run produced, professionally written screenplays that contain no "
        "planted problems. Re-point pass / fail from contradiction false "
        "positives to orphan false positives and spurious high-risk flags.",
    )

    _heading(document, "2.5 API contract tests", level=2)
    _para(
        document,
        "Cover the full request path end to end: upload, scripts, orphans, "
        "orphan-graph, simulate/cut, simulate/edit, draft/delete, draft/apply-"
        "edit, draft/undo, and draft/export — with happy-path and malformed "
        "input cases.",
    )


def _build_unconventional(document: Document) -> None:
    """Add the unconventional testing methods section."""
    _heading(document, "3. Unconventional Testing (high return on investment)")
    _para(
        document,
        "Domain-specific methods that stress story logic and NLP resilience in "
        "ways ordinary unit tests cannot.",
    )

    _heading(document, "3.1 Metamorphic — entity-swap isomorphism", level=2)
    _para(
        document,
        "Rename an entity consistently across a script (for example ALICE to "
        "CHARACTER_X, GUN to OBJECT_Y). The dependency graph and orphan set must "
        "stay unchanged. Any difference exposes name, casing, or state leakage "
        "in the NLP layer. Cheapest, most deterministic, highest-value test for "
        "a structure engine; allow a small tolerance for legitimate name "
        "collisions.",
    )

    _heading(document, "3.2 Metamorphic — scene-permutation locality", level=2)
    _para(
        document,
        "Swap two adjacent scenes that share no entities. The rest of the graph "
        "must remain identical, verifying that a local edit produces only a "
        "local change.",
    )

    _heading(document, "3.3 Synthetic Chekhov's-gun generator", level=2)
    _para(
        document,
        "Programmatically generate three-scene plant -> filler -> payoff scripts "
        "across naming and formatting variants (\"a brass key\", \"the key\", "
        "\"it\"; uppercase, lowercase, buried in dialogue). Assert the setup-"
        "payoff edge and the correct scene-function role resolve every time. "
        "Isolates fuzzy NLP behavior in millisecond tests instead of debugging "
        "it inside full-length PDFs.",
    )

    _heading(document, "3.4 Graceful degradation (garbage-in) curve", level=2)
    _para(
        document,
        "Inject OCR noise, missing punctuation, and lowercased sluglines at 5%, "
        "10%, 25%, and 50%. Orphan counts should drift gradually and the engine "
        "should downgrade from full to limited structure mode rather than crash "
        "or produce wildly unstable output.",
    )

    _heading(document, "3.5 Narrative chaos — single-word sensitivity", level=2)
    _para(
        document,
        "Delete one noun or named entity at a time and measure the change in "
        "graph edges. Use it as an outlier detector: flag hyper-sensitive nodes "
        "(one word removal un-links many scenes) and under-sensitive nodes "
        "(deleting a major action block changes nothing).",
    )

    _heading(document, "3.6 Deferred methods", level=2)
    _bullet(
        document,
        "Human-vs-engine correlation (blind readers ranking vital scenes, "
        "Spearman rho >= 0.75): valuable for trust, but expensive; revisit once "
        "the structure corpus exists.",
    )
    _bullet(
        document,
        "Client memory / DOM-thrashing benchmarks: not applicable until a "
        "browser client UI exists (the current product is a Python engine plus "
        "API).",
    )


def _build_matrix(document: Document) -> None:
    """Add the summary matrix of all methods."""
    _heading(document, "4. Summary Matrix")
    _table(
        document,
        ["Method", "Type", "Primary target", "Value"],
        [
            [
                "Golden-file structure corpus",
                "Conventional",
                "Orphans, cut/edit, functions",
                "Core precision/recall CI gate",
            ],
            [
                "Existing-YAML wiring",
                "Conventional",
                "Orphans, simulate edit",
                "Free coverage already defined",
            ],
            [
                "Unit / regression",
                "Conventional",
                "Parser, graph, orphan logic",
                "Fast per-component safety net",
            ],
            [
                "False-positive corpus",
                "Conventional",
                "Clean produced scripts",
                "Guards against over-flagging",
            ],
            [
                "API contract tests",
                "Conventional",
                "All endpoints",
                "Protects the public surface",
            ],
            [
                "Entity-swap isomorphism",
                "Unconventional",
                "NLP / entity extraction",
                "Guarantees naming neutrality",
            ],
            [
                "Scene-permutation locality",
                "Unconventional",
                "Dependency graph",
                "Confirms edits stay local",
            ],
            [
                "Chekhov generator",
                "Unconventional",
                "Edge-resolution logic",
                "Isolates NLP edge cases fast",
            ],
            [
                "Graceful degradation curve",
                "Unconventional",
                "Ingest resilience",
                "Ensures no catastrophic failure",
            ],
            [
                "Single-word sensitivity",
                "Unconventional",
                "Graph sensitivity",
                "Finds brittle / dead nodes",
            ],
        ],
    )


def _build_next_steps(document: Document) -> None:
    """Add the recommended sequencing section."""
    _heading(document, "5. Recommended Sequence")
    _number(document, "Decouple CI from contradiction detection (done).")
    _number(document, "Archive the contradiction corpus and assets reversibly.")
    _number(
        document,
        "Design the structure ground-truth schema (orphans, edges, cut, edit, "
        "functions) with a reusable template.",
    )
    _number(
        document,
        "Seed 15–20 micro-scripts, reusing the orphan-spec and simulate demos.",
    )
    _number(
        document,
        "Write a structure baseline scorer and set a new CI gate on structure "
        "metrics.",
    )
    _number(
        document,
        "Add the two highest-value metamorphic tests (entity-swap isomorphism "
        "and the Chekhov generator); they need no corpus.",
    )


def _build_plain_english(document: Document) -> None:
    """Add the plain-English 'what needs to be done' section."""
    _heading(document, "6. What Needs To Be Done — In Plain English")
    _para(
        document,
        "Here is the whole plan without the jargon. Right now most of our "
        "automated checks test a feature we are removing (spotting story "
        "contradictions). We need checks that test the features we are keeping: "
        "loose scenes, what breaks when you cut or rewrite a scene, and whether "
        "a setup earlier in the script still pays off later. The steps below get "
        "us there.",
    )

    _heading(document, "Step 1 — Stop failing builds over the old feature (done)", level=2)
    _para(
        document,
        "Our automated pipeline used to block all work if the old contradiction "
        "feature dipped in accuracy. We have switched that off, so the team is "
        "no longer blocked by a feature we are retiring. The old check can still "
        "be run by hand when someone wants it.",
    )

    _heading(document, "Step 2 — Put the old test material aside safely", level=2)
    _para(
        document,
        "Move the old contradiction scripts and answer keys into a clearly "
        "labelled 'legacy' area. We are not deleting anything, just getting it "
        "out of the way so it does not confuse the new work. If we ever bring "
        "the feature back, it is all still there.",
    )

    _heading(document, "Step 3 — Write down the 'right answers' for a few examples", level=2)
    _para(
        document,
        "For a small set of example scripts, agree in a simple checklist what "
        "the correct result should be: which scenes are loose, which later "
        "scenes break if you remove a given scene, and which early setups should "
        "connect to later payoffs. This checklist is what we measure the tool "
        "against.",
    )

    _heading(document, "Step 4 — Build about 15 to 20 small example scripts", level=2)
    _para(
        document,
        "Write short, deliberate example scripts (five to ten scenes each) that "
        "each contain a known situation, and pair every one with its checklist "
        "of right answers. We can reuse the demo scripts we already have as a "
        "starting point. Quality matters more than quantity.",
    )

    _heading(document, "Step 5 — Build a simple scorekeeper", level=2)
    _para(
        document,
        "Create a small program that runs the tool on every example script, "
        "compares what it found against the right answers, and reports a clear "
        "score (how much it got right, and how often it raised a false alarm). "
        "Wire this score into the automated pipeline as the new pass/fail gate.",
    )

    _heading(document, "Step 6 — Add two clever safety checks", level=2)
    _bullet(
        document,
        "Rename test: rename every character and object in a script, run it "
        "again, and confirm the tool's results do not change. If they do, the "
        "tool is unfairly reacting to specific names.",
    )
    _bullet(
        document,
        "Planted-clue test: automatically create tiny scripts where a clue is "
        "introduced early and used later, and confirm the tool always connects "
        "the two, no matter how the clue is worded.",
    )

    _heading(document, "Where to start now", level=2)
    _para(
        document,
        "Step 1 is complete. The best next actions are Step 3 (agree the "
        "checklist format for right answers) and Step 6's rename test, because "
        "neither needs the full example library to exist first and both deliver "
        "value immediately.",
    )


def build_document() -> Document:
    """Assemble the full testing-process document."""
    document = Document()
    _set_styles(document)
    _build_overview(document)
    _build_conventional(document)
    _build_unconventional(document)
    _build_matrix(document)
    _build_next_steps(document)
    _build_plain_english(document)
    return document


def main() -> None:
    """Write the testing-process Word file to the docs folder."""
    document = build_document()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
