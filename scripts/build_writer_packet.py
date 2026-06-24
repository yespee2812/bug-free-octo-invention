"""Build the screenwriter instruction Word document for corpus participants."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

_REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = _REPO_ROOT / "docs" / "writer_materials" / "ScriptLens_Writer_Instructions.docx"


def _set_document_styles(document: Document) -> None:
    """Apply base font settings to the document's Normal style."""
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _add_title(document: Document, text: str) -> None:
    """Add a centered document title."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x11, 0x32, 0x4D)


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    """Add a styled section heading."""
    document.add_heading(text, level=level)


def _add_bullet(document: Document, text: str, bold_prefix: str = "") -> None:
    """Add a bullet paragraph with optional bold prefix."""
    paragraph = document.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text)
    else:
        paragraph.add_run(text)


def _add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Insert a simple table with a header row."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def build_writer_instructions_docx(output_path: Path = OUTPUT_PATH) -> Path:
    """Generate the screenwriter instruction packet as a Word document.

    Args:
        output_path: Destination path for the ``.docx`` file.

    Returns:
        Resolved path to the written document.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _set_document_styles(document)

    _add_title(document, "ScriptLens — Writer Instructions")
    document.add_paragraph(
        "Continuity error injection brief for corpus participants"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

    document.add_paragraph(
        "Replace bracketed placeholders before sending this packet to writers."
    )

    _add_heading(document, "1. Overview", level=1)
    document.add_paragraph(
        "Thank you for agreeing to take part in this project. ScriptLens is a tool "
        "that helps screenwriters catch story continuity problems — timeline slips, "
        "props that change hands with no explanation, characters who contradict "
        "their backstory, and similar logic breaks."
    )
    document.add_paragraph(
        "You will receive three starter screenplays from us. Your job is to inject "
        "deliberate continuity mistakes into each one, then return the edited "
        "scripts plus a written answer sheet documenting every mistake you planted."
    )
    important = document.add_paragraph()
    run = important.add_run("Important: ")
    run.bold = True
    important.add_run(
        "Write the way you normally would on a professional job — natural dialogue, "
        "action, and scene description. Do not write for a machine. We will run "
        "ScriptLens on your scripts and compare the tool's output to your answer "
        "sheet. The answer sheet is the key; we are testing the software, not you."
    )

    _add_heading(document, "2. What you receive from us", level=1)
    _add_table(
        document,
        ["#", "Script", "Scenes", "Your task"],
        [
            ["1", "5-scene script", "5", "Inject 2–3 deliberate continuity errors"],
            ["2", "10-scene script", "10", "Inject 4–5 deliberate continuity errors"],
            [
                "3",
                "Full-length feature script",
                "Full feature",
                "Inject 8–12 deliberate continuity errors",
            ],
        ],
    )
    document.add_paragraph()
    document.add_paragraph("We provide all three scripts. You do not write them from scratch.")
    document.add_paragraph(
        "The 5-scene and 10-scene scripts are short starters we supply. "
        "The full-length script is a pre-produced screenplay we supply. "
        "Keep the original story, characters, and scene headings wherever possible. "
        "Change or add only what you need to plant errors — like a light continuity "
        "pass, not a rewrite."
    )

    _add_heading(document, "3. What you return", level=1)
    document.add_paragraph(
        "For each of the three scripts, send two files:"
    )
    _add_bullet(document, "The edited script (.fountain, .pdf, or .txt)")
    _add_bullet(document, "One Error Injection Log — answer sheet (YAML preferred; Word or Google Doc acceptable)")
    document.add_paragraph("Total: 6 files (3 scripts + 3 logs).")
    document.add_paragraph()
    document.add_paragraph("File naming — use your surname or initials:")
    naming = document.add_paragraph()
    naming.style = "No Spacing"
    for line in [
        "[SURNAME]_5scene_errors.fountain",
        "[SURNAME]_5scene_ERROR_LOG.yaml",
        "[SURNAME]_10scene_errors.fountain",
        "[SURNAME]_10scene_ERROR_LOG.yaml",
        "[SURNAME]_feature_errors.fountain",
        "[SURNAME]_feature_ERROR_LOG.yaml",
    ]:
        document.add_paragraph(line, style="List Bullet")
    document.add_paragraph(
        "You may zip everything as: [SURNAME]_scriptlens_submission.zip"
    )
    document.add_paragraph("Submit to: [YOUR EMAIL / Drive link]")
    document.add_paragraph("Deadline: [DATE]")
    document.add_paragraph("Payment: [AMOUNT / terms]")

    _add_heading(document, "4. Scene numbering", level=1)
    document.add_paragraph(
        "Scene 1 = the first INT. or EXT. heading in the file, top to bottom. "
        "Scene 2 = the second heading, and so on."
    )
    document.add_paragraph(
        "Every error in your log must list an establishing scene (where the first "
        "fact is set up) and a contradicting scene (where the script breaks that fact)."
    )

    _add_heading(document, "5. The 12 types of errors you can inject", level=1)
    document.add_paragraph(
        "Use natural phrasing — dialogue, action, or description. Across all three "
        "scripts combined, use at least 8 different categories."
    )
    _add_table(
        document,
        ["#", "Category", "What the mistake looks like"],
        [
            [
                "1",
                "Character dead then alive",
                "Clearly dead or killed; later active with no valid explanation",
            ],
            [
                "2",
                "Timeline slip",
                "Days, dates, or today/yesterday don't add up in linear order",
            ],
            [
                "3",
                "Role / profession clash",
                "Same person, two incompatible jobs (e.g. surgeon then lawyer)",
            ],
            [
                "4",
                "Prop — wrong owner",
                "Prop with Character A, then Character B — no handoff scene",
            ],
            [
                "5",
                "Prop — destroyed but back",
                "Burned, smashed, or destroyed — then used again",
            ],
            [
                "6",
                "Prop — lost then back",
                "Character loses it; same character has it again with no recovery",
            ],
            [
                "7",
                "Injury — wrong body side",
                "Left arm/leg wound, later same wound on opposite side",
            ],
            [
                "8",
                "Injury — no recovery",
                "Unconscious or seriously hurt, then fine with no hospital or time jump",
            ],
            [
                "9",
                "Relationship — impossible",
                "Same two people: incompatible relations (e.g. siblings + spouses)",
            ],
            [
                "10",
                "Relationship — parent flip",
                "A is B's parent, later B is A's parent",
            ],
            [
                "11",
                "Location clash",
                "Same place: opposite descriptions (abandoned vs busy)",
            ],
            [
                "12",
                "World rule broken (optional)",
                "Rule stated on page, later broken with no explanation",
            ],
        ],
    )
    document.add_paragraph()
    document.add_paragraph("Spread your errors:")
    _add_bullet(document, "5-scene script — across different scenes, not all in Scene 5")
    _add_bullet(document, "10-scene script — opening, middle, and end")
    _add_bullet(document, "Feature script — roughly 2–4 errors per act")

    _add_heading(document, "6. Do NOT log as errors", level=1)
    for item in [
        "Enemies becoming friends, breakups, divorce (valid story arcs)",
        "Fake death explained on the page before the character returns",
        "Flashbacks, dreams, or clearly marked time jumps",
        "Prop handoffs shown on page (gives, hands, steals, finds)",
        "Anything you did not deliberately plant",
    ]:
        _add_bullet(document, item)

    _add_heading(document, "7. How to fill in the answer sheet", level=1)
    document.add_paragraph("One log per script. Copy the attached ERROR_INJECTION_LOG_TEMPLATE.yaml.")
    document.add_paragraph("Header fields:")
    for field in [
        "script_title",
        "filename",
        "writer_name",
        "date",
        'script_type ("5-scene", "10-scene", or "full-length")',
        "total_scenes",
        "base_script_provided_by_us: true (for all scripts we send you)",
    ]:
        _add_bullet(document, field)

    document.add_paragraph()
    document.add_paragraph("Required for each planted error:")
    for field in [
        "error_number",
        "category (plain English from the table above)",
        "establishing_scene and contradicting_scene (integers)",
        "characters_involved and objects_involved (where relevant)",
        "establishing_moment — quote or close paraphrase from the script",
        "contradicting_moment — quote or close paraphrase from the script",
        "how_a_reader_notices — one sentence",
        "writer_intent: deliberate",
    ]:
        _add_bullet(document, field)

    document.add_paragraph()
    document.add_paragraph("Example error entry:")
    example = document.add_paragraph()
    example.style = "No Spacing"
    example.add_run(
        "Category: Timeline slip\n"
        "Establishing scene: 3 — ROSS: Today's Monday. We move at dawn.\n"
        "Contradicting scene: 7 — ROSS: Yesterday was Wednesday.\n"
        "How a reader notices: Monday and Wednesday cannot both be correct."
    ).font.name = "Consolas"

    document.add_paragraph()
    document.add_paragraph("Scene index (required): list every scene heading in order.")
    document.add_paragraph("Summary (required): total_planted_errors and categories_used.")

    _add_heading(document, "8. Checklist before you submit", level=1)
    for item in [
        "All 3 edited scripts returned",
        "All 3 answer sheets returned",
        "Every planted error documented",
        "Nothing accidental in the logs",
        "Scripts read like production drafts",
        'No labels in the script ("CONTINUITY ERROR HERE")',
        "Errors not explained away in the next scene",
        "File names follow the naming convention",
    ]:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run("☐ ").bold = True
        paragraph.add_run(item)

    _add_heading(document, "9. Rights & confidentiality", level=1)
    document.add_paragraph(
        "Your work will be used only for internal testing and improving ScriptLens. "
        "Scripts will not be published, produced, or shared publicly without "
        "separate permission. [Add NDA / agreement reference if applicable.]"
    )

    _add_heading(document, "10. Contact", level=1)
    document.add_paragraph("Name: [Your full name]")
    document.add_paragraph("Project: ScriptLens")
    document.add_paragraph("Email: [Your email]")
    document.add_paragraph("Phone: [Optional]")

    _add_heading(document, "Attachments in this packet", level=1)
    for item in [
        "ScriptLens_Writer_Instructions.docx (this document)",
        "starter_5scene.fountain",
        "starter_10scene.fountain",
        "[Your pre-produced feature script].fountain or .pdf",
        "SCREENWRITER_ERROR_CHEAT_SHEET.pdf",
        "ERROR_INJECTION_LOG_TEMPLATE.yaml",
    ]:
        _add_bullet(document, item)

    document.save(str(output_path))
    return output_path.resolve()


def main() -> None:
    """CLI entry point."""
    written = build_writer_instructions_docx()
    print(f"Wrote Word document: {written}")


if __name__ == "__main__":
    main()
