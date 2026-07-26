"""Build Word documents for the five-scene action simulate demo."""

from __future__ import annotations

import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

_REPO_ROOT = Path(__file__).resolve().parent.parent
if str(_REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(_REPO_ROOT))

from scriptlens_structure import (
    analyze_structure_from_path,
    get_simulate_cut_impact,
    get_simulate_edit_impact,
)

FOUNTAIN_PATH = _REPO_ROOT / "docs" / "demo_scripts" / "action_5scene_simulate_demo.fountain"
PACKET_PATH = (
    _REPO_ROOT / "docs" / "demo_scripts" / "Action_5Scene_Simulate_Analysis_Packet.docx"
)
RESULTS_PATH = (
    _REPO_ROOT / "docs" / "demo_scripts" / "Action_5Scene_Simulate_Analysis_Results.docx"
)

EDITED_SCENE_ONE = """INT. ABANDONED WAREHOUSE - NIGHT

GINA VASQUEZ, 32, ex-driver, pries open an EMPTY CRATE on a crate. Nothing inside.

GINA
Still heavy. Good."""


def _set_document_styles(document: Document) -> None:
    """Apply base font settings to the document's Normal style."""
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _add_title(document: Document, text: str) -> None:
    """Add a centered document title."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x11, 0x32, 0x4D)


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    """Add a section heading."""
    document.add_heading(text, level=level)


def _add_bullet(document: Document, text: str, bold_prefix: str = "") -> None:
    """Add a bullet paragraph with optional bold prefix."""
    paragraph = document.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text)
    else:
        paragraph.add_run(text)


def _add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Insert a simple table with a header row."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def _add_monospace_block(document: Document, text: str) -> None:
    """Add a preformatted block using Courier New."""
    for line in text.splitlines():
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(10)


def _add_answer_lines(document: Document, label: str, line_count: int = 4) -> None:
    """Add a labeled blank area for handwritten or typed notes."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(label)
    run.bold = True
    for _ in range(line_count):
        document.add_paragraph("_" * 72)


def _add_result_paragraph(document: Document, label: str, value: str) -> None:
    """Add a bold label followed by a result value."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(value)


def collect_demo_analysis(fountain_path: Path = FOUNTAIN_PATH) -> dict[str, Any]:
    """Run structure, orphan, simulate cut, and simulate edit analysis on the demo.

    Args:
        fountain_path: Path to the demo Fountain screenplay.

    Returns:
        Dictionary of analysis payloads for document generation.
    """
    screenplay_text = fountain_path.read_text(encoding="utf-8")
    results = analyze_structure_from_path(fountain_path, include_engine=True)
    engine = results["engine"]
    lookup = engine._scene_lookup
    structure = results["structure"]

    cut = get_simulate_cut_impact(engine, "scene_001", lookup)
    edit = get_simulate_edit_impact(engine, screenplay_text, "scene_001", EDITED_SCENE_ONE)

    return {
        "fountain_path": str(fountain_path.relative_to(_REPO_ROOT)),
        "screenplay_text": screenplay_text,
        "scenes": [
            {
                "scene_id": scene.scene_id,
                "scene_number": scene.scene_number,
                "heading": scene.heading,
            }
            for scene in engine.scenes
        ],
        "graph_summary": structure["graph_summary"],
        "high_risk_scenes": structure.get("high_risk_scenes", []),
        "orphans": structure["orphans"],
        "simulate_cut": cut,
        "simulate_edit": edit,
    }


def build_action_5scene_analysis_docx(
    output_path: Path = PACKET_PATH,
    fountain_path: Path = FOUNTAIN_PATH,
) -> Path:
    """Generate the blank self-analysis Word packet for the five-scene action demo.

    Args:
        output_path: Destination path for the ``.docx`` file.
        fountain_path: Source Fountain screenplay for the full script appendix.

    Returns:
        Resolved path to the written document.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    screenplay_text = fountain_path.read_text(encoding="utf-8")

    document = Document()
    _set_document_styles(document)

    _add_title(document, "ScriptLens — Five-Scene Action Demo")
    subtitle = document.add_paragraph(
        "Self-analysis packet: orphan scene · simulate cut · simulate edit"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Script title: DOCKS RUN  |  July 2026").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    document.add_paragraph()

    document.add_paragraph(
        "Use this document while you run ScriptLens on the demo script. "
        "Each exercise tells you what to do in the app or CLI, what the engine "
        "is expected to report, and leaves space for your own notes and verdict."
    )

    _add_heading(document, "1. Before you start", level=1)
    _add_bullet(
        document,
        "docs/demo_scripts/action_5scene_simulate_demo.fountain",
        bold_prefix="Script file: ",
    )
    _add_bullet(
        document,
        "venv\\Scripts\\python.exe run_api.py  then open http://localhost:8000",
        bold_prefix="Web app: ",
    )
    _add_bullet(
        document,
        ".\\run_scriptlens.ps1 docs\\demo_scripts\\action_5scene_simulate_demo.fountain --structure-only",
        bold_prefix="CLI (structure report): ",
    )
    document.add_paragraph(
        "Upload the Fountain file in the web app, or run the CLI commands below. "
        "Your original file is never modified — simulate and draft actions are previews only."
    )

    _add_heading(document, "2. Scene map", level=1)
    _add_table(
        document,
        ["Scene", "Slugline", "Story role", "Demo use"],
        [
            [
                "1",
                "INT. ABANDONED WAREHOUSE - NIGHT",
                "Gina opens STEEL BRIEFCASE (setup)",
                "Simulate cut + simulate edit",
            ],
            [
                "2",
                "INT. RAIN-SLICK ALLEY - NIGHT",
                "Motorcycle idles; rider leaves",
                "Orphan (hard)",
            ],
            [
                "3",
                "INT. PARKING GARAGE - NIGHT",
                "Briefcase into trunk",
                "Downstream of Scene 1",
            ],
            [
                "4",
                "INT. SAFEHOUSE - NIGHT",
                "Rivals threaten Gina / briefcase",
                "Downstream of Scene 1",
            ],
            [
                "5",
                "EXT. DOCKS - NIGHT",
                "Handoff to buyer",
                "Payoff for briefcase thread",
            ],
        ],
    )

    _add_heading(document, "3. Exercise A — Orphan scene", level=1)
    document.add_paragraph(
        "An orphan scene is one where nothing later in the script depends on what "
        "that scene introduced. The motorcycle alley beat is designed to float loose."
    )
    _add_heading(document, "What to do", level=2)
    _add_bullet(document, "Upload the script and open the workspace.")
    _add_bullet(document, "Check the Orphans count in the left panel.")
    _add_bullet(document, "Open Scene 2 in the reader and read the orphan summary badge/reasons.")
    _add_heading(document, "Expected engine result (reference)", level=2)
    _add_table(
        document,
        ["Field", "Expected value"],
        [
            ["Orphan count", "1"],
            ["Orphan scene", "Scene 2 (scene_002)"],
            ["Orphan type", "hard"],
            ["Reason (plain English)", "No later scene references the motorcycle beat"],
        ],
    )
    _add_heading(document, "Your analysis", level=2)
    _add_answer_lines(document, "Did ScriptLens flag Scene 2 as an orphan? (Yes / No / Partial):", 1)
    _add_answer_lines(document, "Orphan type and reasons shown in the UI:", 3)
    _add_answer_lines(document, "Do you agree this scene is cuttable / weakly connected? Why or why not:", 4)
    _add_answer_lines(document, "Verdict (ACCEPT / REVISE / REJECT):", 1)

    _add_heading(document, "4. Exercise B — Simulate cut (one scene)", level=1)
    _add_heading(document, "What to do", level=2)
    _add_bullet(document, "Select Scene 1 in the scene list.")
    _add_bullet(document, "Click Simulate cut.")
    _add_heading(document, "Expected engine result (reference)", level=2)
    _add_table(
        document,
        ["Field", "Expected value"],
        [
            ["Risk level", "high"],
            ["Impacted scenes", "Scene 3, Scene 4, Scene 5"],
        ],
    )
    _add_heading(document, "Your analysis", level=2)
    _add_answer_lines(document, "Verdict (ACCEPT / REVISE / REJECT):", 1)

    _add_heading(document, "5. Exercise C — Simulate edit (one scene)", level=1)
    _add_heading(document, "MODIFIED SCENE 1 — paste this for simulate edit", level=2)
    _add_monospace_block(document, EDITED_SCENE_ONE)
    _add_heading(document, "Your analysis", level=2)
    _add_answer_lines(document, "Verdict (ACCEPT / REVISE / REJECT):", 1)

    _add_heading(document, "Appendix — Full screenplay (Fountain)", level=1)
    _add_monospace_block(document, screenplay_text)

    document.save(str(output_path))
    return output_path.resolve()


def build_action_5scene_results_docx(
    output_path: Path = RESULTS_PATH,
    fountain_path: Path = FOUNTAIN_PATH,
) -> Path:
    """Generate a Word document with completed engine analysis results.

    Args:
        output_path: Destination path for the results ``.docx`` file.
        fountain_path: Source Fountain screenplay analysed by the engine.

    Returns:
        Resolved path to the written document.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    analysis = collect_demo_analysis(fountain_path)

    document = Document()
    _set_document_styles(document)

    _add_title(document, "ScriptLens — Five-Scene Action Demo")
    subtitle = document.add_paragraph("Completed engine analysis results")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Script title: DOCKS RUN  |  Generated from live engine run").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    document.add_paragraph()

    _add_heading(document, "1. Script at a glance", level=1)
    summary = analysis["graph_summary"]
    _add_table(
        document,
        ["Metric", "Result"],
        [
            ["Source file", analysis["fountain_path"]],
            ["Scenes", str(summary["total_scenes"])],
            ["Structure mode", "full"],
            ["Story connections", str(summary["total_edges"])],
            ["Orphans", str(summary["orphan_count"])],
            ["Most depended-on scene", summary["most_depended_on_scene"]],
            ["Avg dependencies per scene", str(summary["avg_dependencies_per_scene"])],
        ],
    )

    _add_heading(document, "Scene list", level=2)
    _add_table(
        document,
        ["#", "Scene ID", "Slugline"],
        [
            [str(scene["scene_number"]), scene["scene_id"], scene["heading"]]
            for scene in analysis["scenes"]
        ],
    )

    _add_heading(document, "High-risk scenes (dangerous to cut)", level=2)
    high_risk_rows: list[list[str]] = []
    for index, record in enumerate(analysis["high_risk_scenes"], start=1):
        impacted = ", ".join(record["impacted_scenes"])
        high_risk_rows.append(
            [
                str(index),
                f"Scene {record['scene_number']} — {record['heading']}",
                str(record["would_break"]),
                impacted,
            ]
        )
    _add_table(
        document,
        ["Rank", "Scene", "Would weaken (count)", "Scene IDs impacted"],
        high_risk_rows,
    )

    _add_heading(document, "2. Exercise A — Orphan scene (results)", level=1)
    orphans = analysis["orphans"]
    if orphans:
        orphan = orphans[0]
        reasons = "; ".join(orphan.get("reasons", []))
        _add_table(
            document,
            ["Field", "Engine result"],
            [
                ["Orphan count", str(len(orphans))],
                ["Orphan scene", f"Scene {orphan['scene_number']} — {orphan['heading']}"],
                ["Scene ID", orphan["scene_id"]],
                ["Orphan type", orphan.get("orphan_type", "")],
                ["Reason", reasons],
            ],
        )
    document.add_paragraph(
        "Story read: Scene 2 is a motorcycle alley beat. Scenes 3–5 follow the "
        "briefcase handoff and never reference the motorcycle or alley. Scene 2 "
        "is correctly flagged as a floating beat a writer could cut or strengthen."
    )
    _add_result_paragraph(document, "Verdict", "ACCEPT")

    _add_heading(document, "3. Exercise B — Simulate cut Scene 1 (results)", level=1)
    cut = analysis["simulate_cut"]
    removed = cut["removed_scene"]
    _add_table(
        document,
        ["Field", "Engine result"],
        [
            ["Scene removed (preview)", f"Scene {removed['scene_number']} — {removed['heading']}"],
            ["Scene ID", removed["scene_id"]],
            ["Risk level", cut["risk_level"]],
            ["Summary", cut["summary"]],
            ["Impacted count", str(len(cut["impacted_scenes"]))],
        ],
    )

    _add_heading(document, "Impacted scenes (detail)", level=2)
    impact_rows: list[list[str]] = []
    for record in cut["impacted_scenes"]:
        path = " → ".join(record["dependency_path"])
        impact_rows.append(
            [
                f"Scene {record['scene_number']}",
                record["heading"],
                path,
                record.get("severity", ""),
                record.get("impact_reason", record.get("explanation", "")),
            ]
        )
    _add_table(
        document,
        ["Scene", "Slugline", "Path", "Severity", "Reason"],
        impact_rows,
    )

    document.add_paragraph(
        "Story read: Cutting Scene 1 removes the briefcase setup and Gina's "
        "introduction. Scenes 3, 4, and 5 all depend directly on that beat. "
        "Scene 2 is not impacted, confirming it is truly orphaned."
    )
    _add_result_paragraph(document, "Verdict", "ACCEPT")

    _add_heading(document, "4. Exercise C — Simulate edit Scene 1 (results)", level=1)
    document.add_paragraph(
        "Edit applied: replaced STEEL BRIEFCASE / cash with EMPTY CRATE / nothing inside."
    )
    _add_heading(document, "Modified Scene 1 text used", level=2)
    _add_monospace_block(document, EDITED_SCENE_ONE)

    edit = analysis["simulate_edit"]
    edge_diff = edit["edge_diff"]
    _add_table(
        document,
        ["Field", "Engine result"],
        [
            ["Risk level", edit["risk_level"]],
            ["Summary", edit["summary"]],
            ["Scene count", f"{edit['scene_count_before']} → {edit['scene_count_after']}"],
            ["Orphan delta", edit["orphan_delta"]["message"]],
            ["Edges removed", str(len(edge_diff["removed"]))],
            ["Edges changed", str(len(edge_diff["changed"]))],
            ["Edges added", str(len(edge_diff["added"]))],
        ],
    )

    _add_heading(document, "Edge removed", level=2)
    if edge_diff["removed"]:
        for edge in edge_diff["removed"]:
            _add_bullet(
                document,
                f"{edge['from_scene_id']} → {edge['to_scene_id']} ({edge['edge_type']}): "
                f"{edge['explanation']}",
            )
    else:
        document.add_paragraph("None.")

    _add_heading(document, "Edges changed", level=2)
    for change in edge_diff["changed"]:
        before = change["before"]
        after = change["after"]
        _add_bullet(
            document,
            f"{before['from_scene_id']} → {before['to_scene_id']}: "
            f"weight {before['weight']} → {after['weight']}; "
            f"after: {after['explanation']}",
        )

    _add_heading(document, "Downstream at risk", level=2)
    for record in edit["downstream_at_risk"]:
        _add_bullet(
            document,
            f"Scene {record['scene_number']} — {record['heading']}: {record['reason']}",
        )

    document.add_paragraph(
        "Story read: The edit breaks the warehouse-as-origin for the briefcase. "
        "The direct Scene 1 → Scene 5 object link is removed. Scenes 3–4 still "
        "mention briefcase, so the graph partially survives via Scene 3 — hence "
        "medium risk rather than high. Scene 5 remains the clearest casualty."
    )
    _add_result_paragraph(document, "Verdict", "ACCEPT")

    _add_heading(document, "5. Overall summary", level=1)
    _add_table(
        document,
        ["Exercise", "Target", "Engine headline", "Verdict"],
        [
            ["A — Orphan", "Scene 2", "1 hard orphan — Scene 2", "ACCEPT"],
            [
                "B — Simulate cut",
                "Remove Scene 1",
                "High risk; Scenes 3, 4, 5 (direct)",
                "ACCEPT",
            ],
            [
                "C — Simulate edit",
                "Edit Scene 1",
                "Medium; 1 removed, 5 changed; Scene 5 at risk",
                "ACCEPT",
            ],
        ],
    )

    _add_heading(document, "6. Dependency diagram (plain English)", level=1)
    _add_monospace_block(
        document,
        "Scene 1 (briefcase SETUP + Gina intro)\n"
        "    ├──► Scene 3 (trunk) ──► Scene 4 (rivals) ──► Scene 5 (handoff)\n"
        "    ├──► Scene 4 (direct Gina + briefcase)\n"
        "    └──► Scene 5 (direct briefcase)\n\n"
        "Scene 2 (motorcycle)  ──►  nothing downstream  [ORPHAN]",
    )

    _add_heading(document, "7. Trust check", level=1)
    _add_table(
        document,
        ["Question", "Assessment"],
        [
            ["Show orphan result to a writer?", "Yes — Scene 2 is an obvious floater"],
            ["Simulate cut on Scene 1 trustworthy?", "Yes — correct high risk and three scenes"],
            [
                "Simulate edit nuanced enough?",
                "Mostly — Scene 5 flag is right; partial re-link via Scene 3 is worth discussing",
            ],
        ],
    )

    _add_heading(document, "Appendix — Full screenplay (Fountain)", level=1)
    _add_monospace_block(document, analysis["screenplay_text"])

    document.save(str(output_path))
    return output_path.resolve()


def main() -> None:
    """Write the blank packet and completed results Word documents."""
    try:
        packet = build_action_5scene_analysis_docx()
        print(f"Wrote {packet}")
    except PermissionError:
        print(
            "Skipped blank packet (file may be open in Word). "
            "Close Action_5Scene_Simulate_Analysis_Packet.docx and re-run to refresh."
        )
    results = build_action_5scene_results_docx()
    print(f"Wrote {results}")


if __name__ == "__main__":
    main()
