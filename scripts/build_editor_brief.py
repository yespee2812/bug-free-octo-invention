"""Build a short professional brief for script editors on ScriptLens."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

_REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = _REPO_ROOT / "docs" / "SCRIPTLENS_EDITOR_BRIEF.docx"


def _set_styles(document: Document) -> None:
    """Apply base font settings."""
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _title(document: Document, text: str) -> None:
    """Add document title."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x11, 0x32, 0x4D)


def _subtitle(document: Document, text: str) -> None:
    """Add document subtitle."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def _heading(document: Document, text: str) -> None:
    """Add a section heading."""
    document.add_heading(text, level=1)


def _para(document: Document, text: str) -> None:
    """Add a body paragraph."""
    document.add_paragraph(text)


def _bullet(document: Document, text: str) -> None:
    """Add a bullet."""
    document.add_paragraph(text, style="List Bullet")


def _table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Add a simple table."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cells[index].text = value
    document.add_paragraph()


def build_document() -> Document:
    """Assemble the editor brief."""
    document = Document()
    _set_styles(document)

    _title(document, "ScriptLens — What the Tool Actually Does")
    _subtitle(
        document,
        "Brief for professional script editors · accurate working description",
    )

    _heading(document, "Purpose")
    _para(
        document,
        "ScriptLens is a structural continuity tool. It tracks whether later "
        "scenes still have the setups they rely on, and which scenes sit "
        "loosely in the story. It does not judge writing quality, theme, "
        "pacing, or emotional craft.",
    )

    _heading(document, "What it reads in each scene")
    _para(document, "From Fountain or cleaned PDF text, it notes:")
    _bullet(document, "Characters (speakers and named people in action)")
    _bullet(document, "Props / objects (especially named or handled items)")
    _bullet(document, "Locations (from scene headings)")
    _bullet(
        document,
        "Clear story facts when stated (ownership, injury/status, "
        "relationships, dates)",
    )
    _bullet(
        document,
        "Dialogue that explicitly looks backward "
        "(e.g. “after what you did,” “since that night”)",
    )

    _heading(document, "1. Orphan scenes")
    _para(
        document,
        "An orphan is a scene that does not meaningfully join the main story "
        "thread.",
    )
    _para(document, "How it decides:")
    _bullet(
        document,
        "It measures ties between scenes using shared characters, shared "
        "locations, shared props/wardrobe, and overall topic similarity.",
    )
    _bullet(
        document,
        "A scene with almost no useful tie into or out of the main thread "
        "is flagged as a hard orphan.",
    )
    _bullet(
        document,
        "A small group of scenes linked to each other but not to the main "
        "thread may be flagged as a loose chain.",
    )
    _bullet(
        document,
        "Opening prologue scenes and montage blocks are usually exempt.",
    )
    _bullet(
        document,
        "If a scene has no earlier tie but clearly feeds later scenes, "
        "it is not treated as a hard orphan.",
    )

    _heading(document, "2. Simulate cut")
    _para(
        document,
        "Preview only. The script is not changed. The question answered: "
        "if this scene is removed, which later scenes lose a setup they "
        "currently rely on?",
    )
    _para(document, "Working mechanism:")
    _bullet(
        document,
        "Earlier scenes are linked to later scenes when people, props, "
        "places, story facts, or backward dialogue are reused.",
    )
    _bullet(
        document,
        "When a scene is cut, ScriptLens finds later scenes that depend "
        "on it through those links.",
    )
    _bullet(
        document,
        "If a later scene can still reach the same setup from an earlier "
        "scene that remains, that later scene is not warned. "
        "Example: cutting a middle “carrier” scene may be safe if Scene 1 "
        "already introduced the same prop or character.",
    )
    _bullet(
        document,
        "Warnings are ranked none / low / medium / high by how many later "
        "scenes remain uniquely dependent on the cut.",
    )

    _heading(document, "3. Simulate edit")
    _para(
        document,
        "Preview only. The question answered: if this scene is rewritten, "
        "which story links appear, disappear, or change?",
    )
    _para(document, "Working mechanism:")
    _bullet(document, "The rewritten scene replaces the original in a temporary copy.")
    _bullet(document, "The whole script is re-read and the links are rebuilt.")
    _bullet(
        document,
        "ScriptLens reports: links removed, links added, links changed, "
        "later scenes that lose a setup from the edited scene, and whether "
        "the orphan count rises or falls.",
    )

    _heading(document, "4. Draft actions (optional)")
    _para(
        document,
        "Separately from preview, a writer can apply a delete or edit to a "
        "working draft, undo, and export Fountain. The original upload "
        "remains unchanged until export.",
    )

    _heading(document, "What it does not do")
    _table(
        document,
        ["Not evaluated", "Examples of alerts we do not produce"],
        [
            [
                "Page / act pacing",
                "“Act 2 is 12% shorter than standard pacing.”",
            ],
            [
                "Emotional arcs",
                "“Character shifts from Joy to Rage with no bridge scene.”",
            ],
            [
                "A / B / C plot tracking",
                "“Romance subplot has no progression for 35 pages.”",
            ],
            [
                "Literary quality",
                "Dialogue polish, theme strength, tone criticism",
            ],
        ],
    )

    _heading(document, "Honest scope statement")
    _para(
        document,
        "ScriptLens currently serves the Information / continuity layer: "
        "setups, reuses, cut impact, edit impact, and loosely attached scenes. "
        "It is closest to a structural continuity pass, not a full editorial "
        "coverage of pacing, character emotion, or subplot progression.",
    )

    return document


def main() -> None:
    """Write the editor brief Word file."""
    document = build_document()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
